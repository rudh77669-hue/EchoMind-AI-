from flask import render_template, request, Flask,redirect,url_for,session,flash,send_file
import joblib
from sklearn.metrics.pairwise import cosine_similarity
from dotenv import load_dotenv
import openai as op
import pandas as pd
import csv
from voice_model import speak
import threading
# subprocess is import due to run train_model.py  again 
import subprocess
from image_to_text import image
from user_detail import chek_credidentals,for_signup,for_update
import datetime
from audio_model import audio
from text_to_image import text_to_image
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from otp_sending import send_otp_email
import io, os, re

messages=[{"role":"system","content":"You are a helpful assistant"}]

model_list = {
    "claude": "anthropic/claude-3-haiku",
    "mistral": "mistralai/mistral-7b-instruct",
    "llama": "meta-llama/llama-3-70b-instruct",
    "gpt": "openai/gpt-3.5-turbo"
}
# Load environment variables
load_dotenv()
api_key = os.getenv("OPEN_ROUTER_API_KEY")
op.api_key = api_key
op.api_base = "https://openrouter.ai/api/v1"

# Flask app
app = Flask(__name__)
app.secret_key=os.getenv("SECRET_KEY")
# Load models and data
tfid_matrix = joblib.load("tfid_matrix.pkl")
vectorizer = joblib.load("vectorizer.pkl")
df = pd.read_csv("local_QA.csv")   

# chat history

user_chats = {}          
current_chat_ids = {} 

def save_chats(user_id, chat_id, new_messages):
    """Save only the new messages for this user & chat."""
    if not new_messages:
        return
    df_chats = [[user_id, chat_id, role, message] for role, message in new_messages]
    
    pd.DataFrame(df_chats, columns=["user_id", "chat_id", "role", "message"]).to_csv(
        "chat_history.csv",
        mode="a",
        header=not os.path.exists("chat_history.csv"),
        index=False,
        quoting=csv.QUOTE_ALL,
    )

def load_chats(user_id):
    """Load all chats for a specific user."""
    user_chats[user_id] = {}
    current_chat_ids[user_id] = 1
    if os.path.exists("chat_history.csv"):
        df_chats = pd.read_csv("chat_history.csv")
        df_chats = df_chats[df_chats["user_id"] == user_id]  # filter only this user
        for _, row in df_chats.iterrows():
            cid = int(row["chat_id"])
            if cid not in user_chats[user_id]:
                user_chats[user_id][cid] = []
            user_chats[user_id][cid].append((row["role"], row["message"]))
        current_chat_ids[user_id] = max(user_chats[user_id].keys()) if user_chats[user_id] else 1
    else:
        user_chats[user_id][1] = []

#------------sign up----------------

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        user_id = request.form.get("user_id")
        password = request.form.get("password")
        email = request.form.get("email")
        df_user=pd.read_csv("user_detail.csv")
        result_id=df_user.loc[df_user["User Id"] == user_id]
        result_mail=df_user.loc[df_user["Email"] == email]
        if not result_id.empty:
            flash("Please select unique Id","error")
        elif not result_mail.empty:
            flash("Email address already exist","error")
        else:
            alpha="!@#$%&*"
            a=len(password)
            if len(password)>=8:
                if password.isupper() or password.isdigit() or password.islower():
                    flash("Inavlid format password should contain at least one\nuppercase, lowercase, numbers, special chars, common patterns ","warning")
                else:
                    if any(ch in alpha for ch in password):
                        flash("valid password")
                        for_signup(user_id, password,email)
                        flash("✅ Signup successful, please login.", "success")
                        speak("Signup successful, please login.")
                        return redirect(url_for("login"))
                    else:
                        flash("Inavlid format password should contain at least one\nuppercase, lowercase, numbers, special chars, common patterns ","warning")
            else:
                flash("Password must be of length 8",'warning')
            

    return render_template("signup.html")


# ------------------------------------Forgot---------------------------
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    email = session.get('email', "")
    otp_sent = session.get('otp_sent', False)
    otp_verified = session.get('otp_verified', False)


    if request.method == 'POST':
        email = request.form.get('email', "")
        otp = request.form.get('otp', "")
        new_pass = request.form.get('new_pass', "")
        cnfm_pass = request.form.get('cnfm_pass', "")

        df_email = pd.read_csv("user_detail.csv")
        result_email = df_email.loc[df_email["Email"] == email]


        if email and not otp and not new_pass:
            if not result_email.empty:
                global o_t_p
                o_t_p = send_otp_email(email)
                flash("OTP sent to your email", "success")
                return render_template("forgot.html", email=email, otp_sent=True)
            else:
                flash("Enter a valid registered email", "warning")
                return render_template("forgot.html", email=email)

        elif otp and not new_pass:
            try:
                if otp == o_t_p:
                    flash("OTP verified successfully", "success")
                    session['otp_verified'] = True
                    return render_template("forgot.html", email=email, otp_verified=True)
                else:
                    flash("Enter a valid OTP", "warning")
                    return render_template("forgot.html", email=email, otp_sent=True)
            except Exception:
                flash("Please check OTP", "error")
                return render_template("forgot.html", email=email, otp_sent=True)
        elif new_pass and cnfm_pass:
            if cnfm_pass==new_pass:
                alpha="!@#$%&*"
                
                if len(cnfm_pass)>=8:
                    if cnfm_pass.isupper() or cnfm_pass.isdigit() or cnfm_pass.islower():
                        flash("Inavlid format password should contain at least one\nuppercase, lowercase, numbers, special chars, common patterns ","warning")
                        return render_template("forgot.html", email=email, otp_verified=True)
                    else:
                        if any(ch in alpha for ch in cnfm_pass):
                            for_update(email,cnfm_pass)
                            flash("Password Updated", "success")
                            speak("update success")
                            session.pop('otp_verified', None)
                            session.pop('otp_sent', None)
                            session.pop('email', None)
                            return redirect(url_for("login"))
                        else:
                            flash("Inavlid format password should contain at least one\nuppercase, lowercase, numbers, special chars, common patterns ","warning")
                            return render_template("forgot.html", email=email, otp_verified=True)
                else:
                    flash("Password must be of length 8","warning")
                    return render_template("forgot.html", email=email, otp_verified=True)
                
        else:
            flash("Password do not match","warning")
            return render_template("forgot.html", email=email, otp_verified=True)
    return render_template("forgot.html", email=email,otp_sent=otp_sent,otp_verified=otp_verified)

 
    
#------------------------------Login-----------------------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        user_id = request.form.get("user_id")
        password = request.form.get("password")

        user_row = chek_credidentals(user_id, password)
        if user_row is not None:  # login success
            session["user_id"] = user_id
            flash("✅ Login success", "success")
            speak("Login success")
            load_chats(user_id)

            # --- update last login in CSV ---
            df = pd.read_csv("user_detail.csv")
            df = df.applymap(lambda x: str(x).strip().strip('"').strip("'"))

            current_time = datetime.datetime.now().strftime("%I:%M %p")
            current_date = datetime.date.today().strftime("%d/%m/%y")

            df.loc[df["User Id"] == user_id, "Last Login"] = f"{current_date} {current_time}"
            df.to_csv("user_detail.csv", index=False, quoting=csv.QUOTE_ALL, escapechar="\\")

            # 🔥 Load only this user's chat history
            load_chats(user_id)

            return redirect(url_for("new_chat"))
        else:
            flash("❌ Invalid User Id or Password", "error")
            speak("Invalid User Id or Password")

    return render_template("login.html")


#-------------------logout----------------

@app.route("/logout")
def logout():
    session.clear()
    flash("You have been logged out.", "info")
    speak("Loged out Successfully!")
    return redirect(url_for("login"))
    
         
# -------------------Chatbot------------------------------------

@app.route('/')
def index():
    if "user_id" not in session:
        return redirect(url_for("login"))
    return redirect(url_for("new_chat"))


@app.route("/chat/<int:chat_id>", methods=["GET", "POST"])

def open_chat(chat_id):
    global vectorizer, tfid_matrix, df
    # new_messages=[]
    if "user_id" not in session:
        return redirect(url_for("login"))
    
    user_id=session["user_id"]
    
    if user_id not in user_chats:
        load_chats(user_id)
    
    if chat_id not in user_chats[user_id]:
        user_chats[user_id][chat_id]=[]
        
    current_chat_ids[user_id]=chat_id
    
    new_messages=[]  
    reply="" 
    action = request.form.get("action", "chat") 
    user_input = request.form.get("user_input", "").strip()
    image_file=request.files.get("image_input")
    audio_file=request.files.get("audio_file")
    try:
#----------------------------------image files------------------------------

        if action == "text_to_image" and user_input:
            messages.append({"role":"user","content":user_input})
            output_path = text_to_image(user_input)   # this should save file under /static/
            user_chats[user_id][chat_id].append(("user", user_input))
            user_chats[user_id][chat_id].append(("bot", f"<img src='/{output_path}' width='400'>"))
            messages.append({"role":"assistant","content":output_path})
            return redirect(url_for("open_chat", chat_id=chat_id))
            
            

        elif action == "image_to_text" and image_file and image_file.filename != "":
            messages.append({"role":"user","content":image_file.filename})
            print(image_file)
            reply = image(image_file)  # use OCR function
            reply=f"{reply} Can I also summarize this?"
            ans_image= str(reply).replace("\r\n", " ").replace("\n", " ").replace("\r", " ")

            user_chats[user_id][chat_id].append(("user", image_file.filename))
            user_chats[user_id][chat_id].append(("assistant", ans_image))
            new_messages.append(("user",image_file.filename))
            new_messages.append(("assistant",ans_image))
            messages.append({"role":"assistant","content":ans_image})
            
#----------------------------------audio files------------------------------

        elif action=="audio_to_text" and audio_file and audio_file.filename!="":
            messages.append({"role":"user","content":audio_file.filename})
            reply=audio(audio_file)
            reply=f"Extracted text: {reply}"
            ans_audio= str(reply).replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
            user_chats[user_id][chat_id].append(("user",audio_file.filename))
            user_chats[user_id][chat_id].append(("assistant", ans_audio))
            new_messages.append(("user",audio_file.filename))
            new_messages.append(("assistant",ans_audio)) 
            messages.append({"role":"assistant","content":ans_audio})
                       
            
        elif action == "chat" and user_input:
            messages.append({"role":"user","content":user_input})
            # Convert user input into vector
            user_vec = vectorizer.transform([user_input])
            similarity = cosine_similarity(user_vec, tfid_matrix).flatten()
            # Find best match
            best_match = similarity.argmax()
            confidence = similarity[best_match]
            print(confidence)
            
        # check which model will reply to user
            if confidence > 0.7:  
                reply = df['Answer'][best_match]
                user_chats[user_id][chat_id].append(("user", user_input))
                user_chats[user_id][chat_id].append(("assistant", reply))
                new_messages.append(("user",user_input))
                new_messages.append(("assistant",reply))  
                messages.append({"role": "assistant", "content": reply})
                            
            else:
                print("fallback to Models")
                model_select = request.form.get('model_choice')
                model_choice = model_list.get(model_select,"meta-llama/llama-3-70b-instruct")
                print(model_choice)
                
                response=op.ChatCompletion.create(
                    model=model_choice,
                    messages=messages
                    # messages=[
                    #     {"role":"system","content":"You are a helpful assistant."},
                    #     {"role":"user","content":user_input}]
                )
                reply=response['choices'][0]['message']['content']
                messages.append({"role": "assistant", "content": reply})
                
                # remove all extra parameters from output to save output in single line
                q = str(user_input).replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
                a = str(reply).replace("\r\n", " ").replace("\n", " ").replace("\r", " ").replace(","," ")
                
                new_entry = pd.DataFrame([[q,a]], columns=["Question", "Answer"])
                
                #add data to local model
                new_entry.to_csv(
                    "local_QA.csv",
                    mode="a",
                    header=False,
                    index=False,
                    quoting=csv.QUOTE_ALL,   # quote the whole field
                    escapechar="\\",         # escape internal quotes if any
                )
                # subprocess.run(["python", "train_model.py"])
                threading.Thread(target=lambda: subprocess.run(["python", "train_model.py"])).start()

                
                 # reload updated models into memory (so this session can use them)

                vectorizer = joblib.load("vectorizer.pkl")
                tfid_matrix = joblib.load("tfid_matrix.pkl")
                df = pd.read_csv("local_QA.csv")

                user_chats[user_id][chat_id].append(("user", user_input))
                user_chats[user_id][chat_id].append(("assistant", reply))
                new_messages.append(("user",user_input))
                new_messages.append(("assistant",reply))
        if new_messages:
            save_chats(user_id, chat_id, new_messages)
             
    except Exception as e:
        reply = f"error: {e}"
        print(reply)
    

    # import datetime
    current_date = datetime.date.today().strftime("%d/%m/%Y")
    
    # Ensure current_chat_id exists
    chat_history = user_chats[user_id][chat_id]

    ren = render_template("index.html", 
                          chat_history=chat_history, 
                          current_date=current_date,
                          chats=user_chats[user_id],
                          current_chat_id=chat_id,
                          user_id=user_id)


    if reply:
        threading.Thread(target=speak, args=(reply,)).start()
    return ren



@app.route("/export_chat/<int:chat_id>/<file_type>")
def export_chat(chat_id, file_type):
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    chat_history = user_chats.get(user_id, {}).get(chat_id, [])

    if not chat_history:
        return "No chat history found for this session.", 404

    # ---------- TEXT EXPORT ----------
    if file_type == "txt":
        output = io.StringIO()
        for role, message in chat_history:
            output.write(f"{role.capitalize()}: {message}\n")
        return send_file(
            io.BytesIO(output.getvalue().encode()),
            as_attachment=True,
            download_name=f"chat_{chat_id}.txt",
            mimetype="text/plain"
        )

    # ---------- PDF EXPORT ----------
    elif file_type == "pdf":
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        elements = []

        def fix_img_tags(text):
            """Ensure <img> tags are self-closed."""
            return re.sub(r"<img([^>/]*)>", r"<img\1/>", text)

        for role, message in chat_history:
            # Check if the message contains an image
            if "<img" in message:
                match = re.search(r"src=['\"]([^'\"]+)['\"]", message)
                if match:
                    img_path = match.group(1).lstrip("/")  # remove leading /
                    if os.path.exists(img_path):
                        elements.append(Paragraph(f"<b>{role.capitalize()}:</b>", styles["Normal"]))
                        elements.append(Image(img_path, width=400, height=300))
                        elements.append(Spacer(1, 12))
                        continue  # Skip text paragraph for this message

            # If no image, treat as text
            safe_message = fix_img_tags(message)
            elements.append(Paragraph(f"<b>{role.capitalize()}:</b> {safe_message}", styles["Normal"]))
            elements.append(Spacer(1, 8))  # spacing between messages

        # Build and return PDF
        doc.build(elements)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"chat_{chat_id}.pdf",
            mimetype="application/pdf"
        )

    # ---------- DOCX EXPORT ----------
    elif file_type == "docx":
        doc = Document()
        for role, message in chat_history:
            if "<img" in message:
                match = re.search(r"src=['\"]([^'\"]+)['\"]", message)
                if match:
                    img_path = match.group(1).lstrip("/")
                    if os.path.exists(img_path):
                        doc.add_paragraph(f"{role.capitalize()}:")
                        doc.add_picture(img_path, width=None)
                        continue
            doc.add_paragraph(f"{role.capitalize()}: {message}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"chat_{chat_id}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ---------- INVALID FILE TYPE ----------
    return "Invalid file type", 400


@app.route("/new_chat")
def new_chat():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    if user_id not in user_chats:
        load_chats(user_id)

    new_id = max(user_chats[user_id].keys(), default=0) + 1
    user_chats[user_id][new_id] = []
    current_chat_ids[user_id] = new_id

    return redirect(url_for("open_chat", chat_id=new_id))


if __name__ == "__main__":
    app.run(debug=True)



